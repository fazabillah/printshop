import io
import shutil
import subprocess
from pathlib import Path
from unittest.mock import patch, MagicMock

import pytest
from docx import Document
from docx.oxml.ns import qn
from lxml import etree

from app.services.template_service import (
    _replace_in_paragraph,
    _render,
    _convert_to_pdf,
    fill_hardbound_docx,
    fill_cd_case_docx,
    TemplateConversionError,
)

REAL_TEMPLATES_DIR = Path(__file__).parent.parent / "templates"


# ─── Helpers ──────────────────────────────────────────────────────────────────

def _make_para_with_split_runs(texts: list[str]) -> etree._Element:
    """Build a <w:p> element with one <w:r><w:t> per text item."""
    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    p = etree.Element(f"{{{W}}}p")
    for text in texts:
        r = etree.SubElement(p, f"{{{W}}}r")
        t = etree.SubElement(r, f"{{{W}}}t")
        t.text = text
    return p


def _para_full_text(p_el) -> str:
    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    return "".join(t.text or "" for t in p_el.iter(f"{{{W}}}t"))


def _make_stub_docx_with_placeholder(tmp_path: Path, placeholder: str) -> Path:
    """Create a minimal DOCX file containing one paragraph with the given placeholder."""
    doc = Document()
    doc.add_paragraph(placeholder)
    path = tmp_path / "stub.docx"
    doc.save(path)
    return path


# ─── Unit tests: _replace_in_paragraph ────────────────────────────────────────

def test_replace_single_run():
    p = _make_para_with_split_runs(["{{FULL_NAME}}"])
    _replace_in_paragraph(p, {"{{FULL_NAME}}": "JAFNI SYAZANI"})
    assert _para_full_text(p) == "JAFNI SYAZANI"


def test_replace_split_across_runs():
    """Placeholder split across two runs must still be replaced."""
    p = _make_para_with_split_runs(["{{FULL_", "NAME}}"])
    _replace_in_paragraph(p, {"{{FULL_NAME}}": "ALI BIN AHMAD"})
    assert _para_full_text(p) == "ALI BIN AHMAD"


def test_replace_multiple_placeholders_in_one_paragraph():
    p = _make_para_with_split_runs(["{{TITLE}} by {{FULL_NAME}}"])
    _replace_in_paragraph(p, {"{{TITLE}}": "THESIS", "{{FULL_NAME}}": "ALI"})
    assert _para_full_text(p) == "THESIS by ALI"


def test_no_placeholder_leaves_paragraph_unchanged():
    p = _make_para_with_split_runs(["UNIVERSITI TEKNOLOGI PETRONAS"])
    _replace_in_paragraph(p, {"{{TITLE}}": "X"})
    assert _para_full_text(p) == "UNIVERSITI TEKNOLOGI PETRONAS"


# ─── Integration: replace in body paragraphs ──────────────────────────────────

def test_fill_paragraph_in_real_docx(tmp_path):
    # Build a stub template in a fake templates dir, then call _render with it
    templates_dir = tmp_path / "templates"
    templates_dir.mkdir()
    stub = _make_stub_docx_with_placeholder(templates_dir, "{{FULL_NAME}}")
    # _make_stub_docx saves to a path derived from the dir; rename to match the expected name
    stub.rename(templates_dir / "test.docx")

    mapping = {"{{FULL_NAME}}": "TEST STUDENT"}

    import app.services.template_service as ts
    original_dir = ts._TEMPLATES_DIR
    ts._TEMPLATES_DIR = templates_dir
    try:
        result_path = _render("test.docx", mapping, tmp_path)
    finally:
        ts._TEMPLATES_DIR = original_dir

    doc = Document(result_path)
    full_text = " ".join(p.text for p in doc.paragraphs if p.text.strip())
    assert "TEST STUDENT" in full_text
    assert "{{FULL_NAME}}" not in full_text


def test_fill_hardbound_with_real_template(tmp_path):
    if not (REAL_TEMPLATES_DIR / "hardbound.docx").exists():
        pytest.skip("hardbound.docx template not present")

    fields = {
        "full_name": "Jafni Syazani bin Mohd Nazrin",
        "thesis_title": "Experimental Investigation of Depth Effects on CO2 Corrosion",
        "student_id": "21000201",
        "course_code": "CV",
        "degree": "B.Eng (Hons) Civil Engineering",
        "year": "JAN 2026",
    }
    docx_path = fill_hardbound_docx(fields, tmp_path)
    assert docx_path.exists()

    doc = Document(docx_path)
    all_text = " ".join(p.text for p in doc.paragraphs)
    assert "{{" not in all_text, "Unfilled placeholder found in hardbound.docx"
    assert "JAFNI SYAZANI" in all_text.upper()
    assert "CIVIL ENGINEERING" in all_text.upper()
    assert "JAN 2026" in all_text.upper()


def test_fill_cd_case_with_real_template(tmp_path):
    if not (REAL_TEMPLATES_DIR / "cd_case.docx").exists():
        pytest.skip("cd_case.docx template not present")

    fields = {
        "full_name": "Jafni Syazani bin Mohd Nazrin",
        "thesis_title": "Experimental Investigation of Depth Effects",
        "student_id": "21000201",
        "course_code": "CV",
        "degree": "B.Eng (Hons) Civil Engineering",
        "year": "JAN 2026",
    }
    docx_path = fill_cd_case_docx(fields, tmp_path)
    assert docx_path.exists()

    from app.services.template_service import _TEMPLATES_DIR
    from lxml import etree as ET
    import zipfile

    with zipfile.ZipFile(docx_path) as z:
        with z.open("word/document.xml") as f:
            xml = f.read().decode("utf-8")

    assert "{{" not in xml, "Unfilled placeholder found in cd_case.docx"
    assert "21000201" in xml


# ─── _convert_to_pdf ──────────────────────────────────────────────────────────

def test_convert_to_pdf_invokes_soffice(tmp_path):
    fake_docx = tmp_path / "test.docx"
    fake_docx.write_bytes(b"fake")
    fake_pdf = tmp_path / "test.pdf"
    fake_pdf.write_bytes(b"%PDF-1.4 fake")  # must exist after mock

    with patch("app.services.template_service.subprocess.run") as mock_run:
        mock_run.return_value = MagicMock(returncode=0, stderr="")
        result = _convert_to_pdf(fake_docx, tmp_path)

    assert mock_run.called
    cmd = mock_run.call_args[0][0]
    assert "--headless" in cmd
    assert "--convert-to" in cmd
    assert "pdf" in cmd
    assert str(fake_docx) in cmd
    assert result == fake_pdf


def test_convert_to_pdf_raises_on_nonzero_exit(tmp_path):
    fake_docx = tmp_path / "test.docx"
    fake_docx.write_bytes(b"fake")

    with patch("app.services.template_service.subprocess.run") as mock_run:
        mock_run.return_value = MagicMock(returncode=1, stderr="error from soffice")
        with pytest.raises(TemplateConversionError, match="exited 1"):
            _convert_to_pdf(fake_docx, tmp_path)


def test_convert_to_pdf_raises_when_binary_missing(tmp_path):
    fake_docx = tmp_path / "test.docx"
    fake_docx.write_bytes(b"fake")

    with patch("app.services.template_service.subprocess.run", side_effect=FileNotFoundError):
        with pytest.raises(TemplateConversionError, match="not found"):
            _convert_to_pdf(fake_docx, tmp_path)


def test_convert_to_pdf_raises_on_timeout(tmp_path):
    fake_docx = tmp_path / "test.docx"
    fake_docx.write_bytes(b"fake")

    with patch("app.services.template_service.subprocess.run", side_effect=subprocess.TimeoutExpired(cmd="soffice", timeout=60)):
        with pytest.raises(TemplateConversionError, match="timed out"):
            _convert_to_pdf(fake_docx, tmp_path)
